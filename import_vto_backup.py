#!/usr/bin/env python3
"""Import VTO data from all three divisions into multi-tenant database"""

import sqlite3
from docx import Document
from datetime import datetime
from pathlib import Path

DATABASE_PATH = Path(__file__).parent / 'eos_data.db'
DATASHEETS_PATH = Path(__file__).parent / 'datasheets'

def parse_vto_document(docx_path):
    """Extract VTO data from the Word document"""
    doc = Document(docx_path)
    
    # Parse Table 1 (Main VTO)
    table1 = doc.tables[0] if doc.tables else None
    if not table1:
        return None
    
    # Extract Core Values (Row 0, Col 2)
    core_values_text = table1.rows[0].cells[2].text.strip()
    core_values = [v.strip() for v in core_values_text.split('\n') if v.strip()]
    
    # Extract Core Focus (Row 2, Col 2)
    core_focus_text = table1.rows[2].cells[2].text.strip()
    passion = niche = cash_flow = ""
    for line in core_focus_text.split('\n'):
        if 'Passion:' in line:
            passion = line.split('Passion:')[1].strip()
        elif 'Niche:' in line or 'Our Niche:' in line:
            niche = line.split('Niche:')[1].strip()
        elif 'Cash Flow' in line:
            cash_flow = line.split('Driver:')[1].strip() if 'Driver:' in line else line
    
    # Extract Core Target (Row 3, Col 2)
    core_target_text = table1.rows[3].cells[2].text.strip()
    
    # Extract Marketing Strategy (Row 4, Col 2)
    marketing_text = table1.rows[4].cells[2].text.strip()
    uniques = guarantee = proven_process = target_market = ""
    current_section = None
    for line in marketing_text.split('\n'):
        if 'Uniques:' in line:
            current_section = 'uniques'
            uniques = line.split('Uniques:')[1].strip() if ':' in line else ""
        elif 'Guarantee:' in line:
            current_section = 'guarantee'
            guarantee = line.split('Guarantee:')[1].strip()
        elif 'Proven Process:' in line:
            current_section = 'proven_process'
            proven_process = line.split('Proven Process:')[1].strip()
        elif 'Target Market' in line or '"The List"' in line:
            current_section = 'target_market'
            target_market = line.split(':')[1].strip() if ':' in line else ""
        elif line.strip() and current_section:
            if current_section == 'uniques':
                uniques += "\n" + line.strip()
            elif current_section == 'target_market':
                target_market += " " + line.strip()
    
    # Extract 3-Year Picture (Row 1, Col 3)
    three_year_text = table1.rows[1].cells[3].text.strip()
    future_date = revenue = profit = measurables = vision = ""
    current_section = None
    for line in three_year_text.split('\n'):
        if 'Future Date:' in line:
            future_date = line.split('Future Date:')[1].strip()
        elif 'Revenue:' in line:
            revenue = line.split('Revenue:')[1].strip()
        elif 'Profit:' in line:
            profit = line.split('Profit:')[1].strip()
        elif 'Measurables:' in line:
            measurables = line.split('Measurables:')[1].strip()
        elif 'What does it look like?' in line:
            current_section = 'vision'
            vision = ""
        elif current_section == 'vision' and line.strip():
            vision += line.strip() + "\n"
    
    # Parse Table 2 (1-Year Plan)
    table2 = doc.tables[1]
    one_year_text = table2.rows[1].cells[0].text.strip()
    year_date = year_revenue = year_profit = year_measurables = year_goals = ""
    current_section = None
    for line in one_year_text.split('\n'):
        if 'Future Date:' in line:
            year_date = line.split('Future Date:')[1].strip()
        elif 'Revenue:' in line:
            year_revenue = line.split('Revenue:')[1].strip()
        elif 'Profit:' in line:
            year_profit = line.split('Profit:')[1].strip()
        elif 'Measurables:' in line:
            year_measurables = line.split('Measurables:')[1].strip()
        elif 'Goals for the Year:' in line:
            current_section = 'goals'
            year_goals = ""
        elif current_section == 'goals' and line.strip() and 'press Tab' not in line:
            year_goals += line.strip() + "\n"
    
    return {
        'core_values': core_values,
        'core_focus': {'passion': passion, 'niche': niche, 'cash_flow_driver': cash_flow},
        'core_target': core_target_text,
        'marketing_strategy': {
            'uniques': uniques.strip(),
            'guarantee': guarantee.strip(),
            'proven_process': proven_process.strip(),
            'target_market': target_market.strip()
        },
        'three_year_picture': {
            'future_date': future_date,
            'revenue': revenue,
            'profit': profit,
            'measurables': measurables,
            'vision': vision.strip()
        },
        'one_year_plan': {
            'future_date': year_date,
            'revenue': year_revenue,
            'profit': year_profit,
            'measurables': year_measurables,
            'goals': year_goals.strip()
        }
    }

def import_vto_to_db(db_path, docx_path):
    """Import VTO data into database"""
    vto_data = parse_vto_document(docx_path)
    
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    
    # Create tables from schema
    with open('vto_schema.sql', 'r') as f:
        schema = f.read()
        cursor.executescript(schema)
    
    # Insert Core Values
    for i, value in enumerate(vto_data['core_values']):
        cursor.execute("""
            INSERT INTO vto_core_values (value_text, sort_order, updated_by)
            VALUES (?, ?, ?)
        """, (value, i, 'system_import'))
    
    # Insert Core Focus
    cf = vto_data['core_focus']
    cursor.execute("""
        INSERT INTO vto_core_focus (passion, niche, cash_flow_driver, updated_by)
        VALUES (?, ?, ?, ?)
    """, (cf['passion'], cf['niche'], cf['cash_flow_driver'], 'system_import'))
    
    # Insert Core Target
    cursor.execute("""
        INSERT INTO vto_core_target (target_text, target_date, updated_by)
        VALUES (?, ?, ?)
    """, (vto_data['core_target'], '2035-12-31', 'system_import'))
    
    # Insert Marketing Strategy
    ms = vto_data['marketing_strategy']
    cursor.execute("""
        INSERT INTO vto_marketing_strategy (uniques, guarantee, proven_process, target_market, updated_by)
        VALUES (?, ?, ?, ?, ?)
    """, (ms['uniques'], ms['guarantee'], ms['proven_process'], ms['target_market'], 'system_import'))
    
    # Insert 3-Year Picture
    typ = vto_data['three_year_picture']
    cursor.execute("""
        INSERT INTO vto_three_year_picture (future_date, revenue, profit, measurables, what_does_it_look_like, updated_by)
        VALUES (?, ?, ?, ?, ?, ?)
    """, ('2028-12-31', typ['revenue'], typ['profit'], typ['measurables'], typ['vision'], 'system_import'))
    
    # Insert 1-Year Plan
    oyp = vto_data['one_year_plan']
    cursor.execute("""
        INSERT INTO vto_one_year_plan (future_date, revenue, profit, measurables, goals, updated_by)
        VALUES (?, ?, ?, ?, ?, ?)
    """, ('2026-12-31', oyp['revenue'], oyp['profit'], oyp['measurables'], oyp['goals'], 'system_import'))
    
    conn.commit()
    
    # Print summary
    print("\n✅ VTO IMPORT COMPLETE\n")
    print(f"Core Values: {len(vto_data['core_values'])} imported")
    print(f"Core Focus: passion, niche, cash flow driver")
    print(f"Core Target: {vto_data['core_target'][:60]}...")
    print(f"Marketing Strategy: 4 sections imported")
    print(f"3-Year Picture: Target 12/31/2028 - {typ['revenue']}")
    print(f"1-Year Plan: Target 12/31/2026 - {oyp['revenue']}")
    
    conn.close()

if __name__ == '__main__':
    import_vto_to_db('eos_data.db', 'datasheets/VTO 2026 Q1 Plainwell Site- Official.docx')
