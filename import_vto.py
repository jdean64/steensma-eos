#!/usr/bin/env python3
"""
Import VTO data from all three divisions (Plainwell, Generator, Kalamazoo)
Ensures universal structure across all divisions
"""

import sqlite3
from docx import Document
from datetime import datetime
from pathlib import Path
import json

DATABASE_PATH = Path(__file__).parent / 'eos_data.db'
DATASHEETS_PATH = Path(__file__).parent / 'datasheets'

# Division mapping
DIVISIONS = {
    'Plainwell': {
        'id': 1,
        'file': 'VTO 2026 Q1 Plainwell Site- Official.docx',
        'display_name': 'Steensma Plainwell'
    },
    'Kalamazoo': {
        'id': 2,
        'file': 'VTO 2026 Q1 Kalamazoo Site.docx',
        'display_name': 'Steensma Kalamazoo'
    },
    'Generator': {
        'id': 3,
        'file': 'VTO 2026 Q1 Generator Division.docx',
        'display_name': 'Steensma Generator'
    }
}

def get_db():
    """Get database connection"""
    conn = sqlite3.connect(DATABASE_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def parse_vto_section(text, keywords):
    """Extract data from section based on keywords"""
    result = {}
    current_key = None
    buffer = []
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        # Check if line starts with a keyword
        found_key = None
        for key in keywords:
            if line.upper().startswith(key.upper() + ':') or line.upper().startswith(key.upper()):
                found_key = key
                break
        
        if found_key:
            # Save previous buffer
            if current_key and buffer:
                result[current_key] = ' '.join(buffer).strip()
            # Start new key
            current_key = found_key
            # Get content after keyword
            if ':' in line:
                content = line.split(':', 1)[1].strip()
                buffer = [content] if content else []
            else:
                buffer = []
        elif current_key:
            # Add to current buffer
            buffer.append(line)
    
    # Save last buffer
    if current_key and buffer:
        result[current_key] = ' '.join(buffer).strip()
    
    return result

def parse_vto_document(docx_path):
    """Extract VTO data from Word document with universal structure"""
    try:
        doc = Document(docx_path)
        
        if len(doc.tables) < 2:
            print(f"  ⚠️  Document has only {len(doc.tables)} table(s), expected at least 2")
            return None
        
        # Parse Table 1 (Main VTO Components)
        table1 = doc.tables[0]
        
        vto_data = {
            'core_values': [],
            'core_focus': {},
            'ten_year_target': '',
            'marketing_strategy': {},
            'three_year_picture': {},
            'one_year_plan': {}
        }
        
        # Extract Core Values (usually Row 0, Col 2)
        try:
            core_values_text = table1.rows[0].cells[2].text.strip()
            vto_data['core_values'] = [v.strip() for v in core_values_text.split('\n') if v.strip() and 'press Tab' not in v.lower()]
        except:
            pass
        
        # Extract Core Focus (usually Row 2, Col 2)
        try:
            core_focus_text = table1.rows[2].cells[2].text.strip()
            vto_data['core_focus'] = parse_vto_section(core_focus_text, ['Passion', 'Niche', 'Cash Flow Driver', 'Our Niche'])
        except:
            pass
        
        # Extract 10-Year Target (Row 3, Col 2)
        try:
            vto_data['ten_year_target'] = table1.rows[3].cells[2].text.strip()
        except:
            pass
        
        # Extract Marketing Strategy (Row 4, Col 2)
        try:
            marketing_text = table1.rows[4].cells[2].text.strip()
            vto_data['marketing_strategy'] = parse_vto_section(marketing_text, 
                ['Uniques', 'Guarantee', 'Proven Process', 'Target Market', '"The List"'])
        except:
            pass
        
        # Extract 3-Year Picture (Row 1, Col 3)
        try:
            three_year_text = table1.rows[1].cells[3].text.strip()
            parsed_3year = parse_vto_section(three_year_text,
                ['Future Date', 'Revenue', 'Profit', 'Measurables', 'What does it look like'])
            
            # Find vision text (everything after "What does it look like?")
            vision_parts = []
            in_vision = False
            for line in three_year_text.split('\n'):
                if 'what does it look like' in line.lower():
                    in_vision = True
                    continue
                if in_vision and line.strip():
                    vision_parts.append(line.strip())
            
            vto_data['three_year_picture'] = {
                'future_date': parsed_3year.get('Future Date', ''),
                'revenue': parsed_3year.get('Revenue', ''),
                'profit': parsed_3year.get('Profit', ''),
                'measurables': parsed_3year.get('Measurables', ''),
                'vision': '\n'.join(vision_parts)
            }
        except Exception as e:
            print(f"  ⚠️  Error parsing 3-year picture: {e}")
        
        # Parse Table 2 (1-Year Plan)
        try:
            table2 = doc.tables[1]
            one_year_text = table2.rows[1].cells[0].text.strip()
            parsed_1year = parse_vto_section(one_year_text,
                ['Future Date', 'Revenue', 'Profit', 'Measurables', 'Goals for the Year'])
            
            # Find goals text
            goals_parts = []
            in_goals = False
            for line in one_year_text.split('\n'):
                if 'goals for the year' in line.lower():
                    in_goals = True
                    continue
                if in_goals and line.strip() and 'press Tab' not in line.lower():
                    goals_parts.append(line.strip())
            
            vto_data['one_year_plan'] = {
                'future_date': parsed_1year.get('Future Date', ''),
                'revenue': parsed_1year.get('Revenue', ''),
                'profit': parsed_1year.get('Profit', ''),
                'measurables': parsed_1year.get('Measurables', ''),
                'goals': '\n'.join(goals_parts)
            }
        except Exception as e:
            print(f"  ⚠️  Error parsing 1-year plan: {e}")
        
        return vto_data
        
    except Exception as e:
        print(f"  ❌ Error parsing document: {e}")
        return None

def import_division_vto(division_name):
    """Import VTO for a specific division"""
    
    if division_name not in DIVISIONS:
        print(f"❌ Unknown division: {division_name}")
        return False
    
    division = DIVISIONS[division_name]
    division_id = division['id']
    doc_file = DATASHEETS_PATH / division['file']
    
    print(f"\n{'='*70}")
    print(f"Importing VTO: {division['display_name']} (ID: {division_id})")
    print(f"Source: {doc_file.name}")
    print(f"{'='*70}\n")
    
    if not doc_file.exists():
        print(f"  ❌ File not found: {doc_file}")
        return False
    
    # Parse the document
    print("  📄 Parsing document...")
    vto_data = parse_vto_document(doc_file)
    
    if not vto_data:
        print("  ❌ Failed to parse VTO data")
        return False
    
    # Connect to database
    conn = get_db()
    cursor = conn.cursor()
    
    # Get organization_id
    cursor.execute("SELECT organization_id FROM divisions WHERE id = ?", (division_id,))
    row = cursor.fetchone()
    if not row:
        print(f"  ❌ Division {division_id} not found in database")
        conn.close()
        return False
    
    organization_id = row['organization_id']
    
    # VTO table already exists in the database with proper schema
    
    # Deactivate old VTO records
    cursor.execute("""
        UPDATE vto 
        SET is_active = 0, updated_at = datetime('now')
        WHERE division_id = ? AND is_active = 1
    """, (division_id,))
    
    deactivated = cursor.rowcount
    if deactivated > 0:
        print(f"  ✓ Deactivated {deactivated} previous VTO record(s)")
    
    # Prepare data for insert
    core_values_json = json.dumps(vto_data.get('core_values', []))
    three_year = vto_data.get('three_year_picture', {})
    one_year = vto_data.get('one_year_plan', {})
    core_focus = vto_data.get('core_focus', {})
    marketing = vto_data.get('marketing_strategy', {})
    
    # Insert new VTO record using existing schema column names
    cursor.execute("""
        INSERT INTO vto (
            organization_id, division_id,
            core_values, 
            core_purpose, core_niche,
            ten_year_target,
            target_market, unique_value_proposition, proven_process, guarantee,
            three_year_revenue, three_year_profit, three_year_measurables,
            one_year_revenue, one_year_profit, one_year_goals,
            effective_date, updated_by
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        organization_id, division_id,
        core_values_json,
        core_focus.get('Passion', ''),
        core_focus.get('Niche', '') or core_focus.get('Our Niche', ''),
        vto_data.get('ten_year_target', ''),
        marketing.get('Target Market', '') or marketing.get('"The List"', ''),
        marketing.get('Uniques', ''),
        marketing.get('Proven Process', ''),
        marketing.get('Guarantee', ''),
        three_year.get('revenue', ''),
        three_year.get('profit', ''),
        three_year.get('measurables', ''),
        one_year.get('revenue', ''),
        one_year.get('profit', ''),
        one_year.get('goals', ''),
        '2026-01-01',  # Q1 2026
        1  # updated by user_id=1
    ))
    
    vto_id = cursor.lastrowid
    print(f"  ✓ Created VTO record (ID: {vto_id})")
    
    # Print summary
    print(f"\n  📊 VTO Summary:")
    print(f"     Core Values: {len(vto_data.get('core_values', []))} items")
    print(f"     Core Focus: {len(vto_data.get('core_focus', {}))} components")
    print(f"     10-Year Target: {len(vto_data.get('ten_year_target', ''))} chars")
    print(f"     Marketing Strategy: {len(vto_data.get('marketing_strategy', {}))} sections")
    print(f"     3-Year Picture: {len(vto_data.get('three_year_picture', {}))} components")
    print(f"     1-Year Plan: {len(vto_data.get('one_year_plan', {}))} components")
    
    conn.commit()
    conn.close()
    
    print(f"\n  ✅ {division_name} VTO imported successfully")
    print(f"{'='*70}\n")
    
    return True

def show_vto_summary():
    """Show VTO status for all divisions"""
    conn = get_db()
    cursor = conn.cursor()
    
    print("\n" + "="*70)
    print("VTO DATA STATUS - ALL DIVISIONS")
    print("="*70)
    
    cursor.execute("""
        SELECT 
            d.id,
            d.name,
            d.display_name,
            v.effective_date,
            v.created_at,
            LENGTH(v.core_values) as cv_len,
            LENGTH(v.core_purpose) as purpose_len,
            LENGTH(v.ten_year_target) as target_len,
            LENGTH(v.unique_value_proposition) as uvp_len,
            LENGTH(v.three_year_measurables) as three_len,
            LENGTH(v.one_year_goals) as one_len
        FROM divisions d
        LEFT JOIN vto v ON d.id = v.division_id AND v.is_active = 1
        WHERE d.is_active = 1
        ORDER BY d.id
    """)
    
    for row in cursor.fetchall():
        print(f"\n{row['display_name']} (ID: {row['id']})")
        if row['effective_date']:
            print(f"  📅 Effective Date: {row['effective_date']}")
            print(f"  📊 Data Size:")
            print(f"     Core Values: {row['cv_len']} chars")
            print(f"     Core Purpose: {row['purpose_len']} chars")
            print(f"     10-Year Target: {row['target_len']} chars")
            print(f"     Unique Value Prop: {row['uvp_len']} chars")
            print(f"     3-Year Picture: {row['three_len']} chars")
            print(f"     1-Year Plan: {row['one_len']} chars")
            print(f"  🕐 Last Updated: {row['created_at']}")
            print(f"  ✅ Status: ACTIVE")
        else:
            print(f"  ⚠️  Status: NO VTO DATA")
    
    print("\n" + "="*70 + "\n")
    conn.close()

def main():
    """Import VTO for all three divisions"""
    print("\n" + "="*70)
    print("VTO IMPORT - UNIVERSAL STRUCTURE")
    print("Importing all three divisions: Plainwell, Kalamazoo, Generator")
    print("="*70 + "\n")
    
    success_count = 0
    
    # Import all three divisions
    for division_name in ['Plainwell', 'Kalamazoo', 'Generator']:
        if import_division_vto(division_name):
            success_count += 1
    
    # Show summary
    show_vto_summary()
    
    print(f"\n✅ Import Complete: {success_count}/3 divisions imported successfully\n")

if __name__ == '__main__':
    main()
